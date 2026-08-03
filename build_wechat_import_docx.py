from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).parent
SOURCE = ROOT / "我用Codex做了一个会说话的个人网站_公众号长文.md"
OUTPUT = ROOT / "个人网页制作教程_公众号导入版.docx"

# WeChat import override: one Unicode family avoids LibreOffice dropping CJK glyphs.
LATIN_FONT = "Arial Unicode MS"
CJK_FONT = "Arial Unicode MS"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "202124"
MUTED = "667085"
LIGHT_FILL = "F4F6F9"
BORDER = "D5DCE5"


def set_font(run, size=None, color=None, bold=None, italic=None, name=LATIN_FONT):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color=INK, bold=False, italic=False):
    style.font.name = LATIN_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style.font.italic = italic
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)


def set_keep(paragraph, with_next=False, together=False):
    paragraph.paragraph_format.keep_with_next = with_next
    paragraph.paragraph_format.keep_together = together
    ppr = paragraph._p.get_or_add_pPr()
    widow = ppr.find(qn("w:widowControl"))
    if widow is None:
        widow = OxmlElement("w:widowControl")
        ppr.append(widow)


def add_left_border(paragraph, color=BLUE, size=18, space=8):
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    borders.append(left)


def add_shading(paragraph, fill=LIGHT_FILL):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading1 = doc.styles["Heading 1"]
    set_style_font(heading1, 16, BLUE, bold=True)
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(10)
    heading1.paragraph_format.line_spacing = 1.1
    heading1.paragraph_format.keep_with_next = True

    heading2 = doc.styles["Heading 2"]
    set_style_font(heading2, 13, BLUE, bold=True)
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)
    heading2.paragraph_format.line_spacing = 1.1
    heading2.paragraph_format.keep_with_next = True

    heading3 = doc.styles["Heading 3"]
    set_style_font(heading3, 12, DARK_BLUE, bold=True)
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(4)
    heading3.paragraph_format.line_spacing = 1.1
    heading3.paragraph_format.keep_with_next = True

    title = doc.styles.add_style("Article Title", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(title, 24, INK, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.08
    title.paragraph_format.keep_with_next = True

    kicker = doc.styles.add_style("Article Kicker", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(kicker, 9, BLUE, bold=True)
    kicker.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(6)
    kicker.paragraph_format.keep_with_next = True

    instruction = doc.styles.add_style("Import Instruction", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(instruction, 9, MUTED)
    instruction.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    instruction.paragraph_format.space_before = Pt(0)
    instruction.paragraph_format.space_after = Pt(18)
    instruction.paragraph_format.line_spacing = 1.15

    caption = doc.styles.add_style("Article Caption", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(caption, 8.5, MUTED, italic=True)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(12)
    caption.paragraph_format.line_spacing = 1.1

    quote = doc.styles.add_style("Article Quote", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(quote, 10.5, DARK_BLUE)
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.right_indent = Inches(0.12)
    quote.paragraph_format.space_before = Pt(8)
    quote.paragraph_format.space_after = Pt(10)
    quote.paragraph_format.line_spacing = 1.208

    source = doc.styles.add_style("Article Source", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(source, 9.5, MUTED)
    source.paragraph_format.space_before = Pt(0)
    source.paragraph_format.space_after = Pt(4)
    source.paragraph_format.line_spacing = 1.208


def add_custom_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)

    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "279")
    ppr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)

    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    rpr.append(fonts)
    level.append(rpr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.append(num_pr)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    rpr.extend((fonts, color, underline))
    run.append(rpr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_text_with_urls(paragraph, text):
    cursor = 0
    for match in re.finditer(r"https?://[^\s，]+", text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_font(run)
        url = match.group(0)
        add_hyperlink(paragraph, url, url)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run)


def set_image_alt(paragraph, alt_text):
    doc_pr = paragraph._p.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("descr", alt_text)
        doc_pr[0].set("title", alt_text)


def image_width(path):
    with Image.open(path) as image:
        width, height = image.size
    ratio = width / height
    if ratio < 0.75:
        return Inches(4.0)
    if ratio < 0.95:
        return Inches(4.5)
    return Inches(6.25)


def add_image(doc, path, alt):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=image_width(path))
    set_image_alt(paragraph, alt)
    return paragraph


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_markdown_table(doc, rows):
    doc.add_page_break()
    headers = [item.strip() for item in rows[0].strip("|").split("|")]
    body = [
        [item.strip() for item in row.strip("|").split("|")]
        for row in rows[2:]
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = 0
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shade_cell(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_font(run, size=9.5, color=DARK_BLUE, bold=True)
    for values in body:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index < 3 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            for run in paragraph.runs:
                set_font(run, size=9.5)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
        if row_index == 0:
            tr_pr.append(OxmlElement("w:tblHeader"))
    set_table_geometry(table, [2160, 1440, 1440, 4320])
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    setup_styles(doc)
    bullet_num_id = add_custom_bullet_numbering(doc)

    core = doc.core_properties
    core.title = "个人网页制作教程，公众号导入版"
    core.subject = "微信公众号图文导入稿"
    core.author = ""
    core.last_modified_by = ""
    core.keywords = ""
    core.comments = ""

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    kicker = doc.add_paragraph(style="Article Kicker")
    kicker.add_run("WECHAT ARTICLE / IMPORT COPY")
    title = doc.add_paragraph(style="Article Title")
    title.add_run("我用 Codex 做了一个会说话的个人网站").add_break()
    title.add_run("保姆级教程，从 0 到正式上线")
    instruction = doc.add_paragraph(style="Import Instruction")
    instruction.add_run(
        "使用提示，标题请单独复制到公众号标题栏，正文从下一段开始复制。"
        "GIF 如果在 Word 中只显示首帧，请在微信编辑器中用原 GIF 原位替换。"
    )

    index = 1
    while index < len(lines):
        line_text = lines[index].strip()
        if not line_text:
            index += 1
            continue

        if line_text.startswith("|"):
            table_rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(lines[index].strip())
                index += 1
            add_markdown_table(doc, table_rows)
            continue

        image_match = re.fullmatch(r"!\[([^]]+)\]\(([^)]+)\)", line_text)
        if image_match:
            alt, relative = image_match.groups()
            add_image(doc, ROOT / relative, alt)
            index += 1
            continue

        if re.fullmatch(r"\[(?:图片|动图) \d+，.+\]", line_text):
            caption = doc.add_paragraph(style="Article Caption")
            caption.add_run(line_text[1:-1])
            index += 1
            continue

        if re.match(r"^\d+、", line_text):
            paragraph = doc.add_paragraph(style="Heading 1")
            paragraph.add_run(line_text)
            set_keep(paragraph, with_next=True)
            index += 1
            continue

        if line_text == "资料入口":
            paragraph = doc.add_paragraph(style="Heading 2")
            paragraph.add_run(line_text)
            set_keep(paragraph, with_next=True)
            index += 1
            continue

        if line_text.startswith(">"):
            quote_text = line_text[1:].strip()
            paragraph = doc.add_paragraph(style="Article Quote")
            add_left_border(paragraph)
            add_shading(paragraph)
            add_text_with_urls(paragraph, quote_text)
            set_keep(paragraph, together=False)
            index += 1
            continue

        if line_text.startswith("- "):
            paragraph = doc.add_paragraph(style="Article Source")
            apply_numbering(paragraph, bullet_num_id)
            add_text_with_urls(paragraph, line_text[2:])
            index += 1
            continue

        paragraph = doc.add_paragraph(style="Normal")
        add_text_with_urls(paragraph, line_text)
        set_keep(paragraph, together=False)
        index += 1

    settings = doc.settings._element
    compress = OxmlElement("w:doNotAutoCompressPictures")
    settings.append(compress)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
